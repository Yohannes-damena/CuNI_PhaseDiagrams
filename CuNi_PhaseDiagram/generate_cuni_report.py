# coding: utf-8
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

def add_heading(doc: Document, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Calibri'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
    return h

def add_paragraph(doc: Document, text: str):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
    return p

def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(text, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
    return p

def add_image(doc: Document, path: str, caption: str, width_inches: float = 6.5):
    if os.path.isfile(path):
        doc.add_picture(path, width=Inches(width_inches))
        last_par = doc.paragraphs[-1]
        last_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic = True
    else:
        add_paragraph(doc, f"[Image missing: {path}]")

def main():
    # Inputs from your latest run (Cu-40Ni)
    solidus_c = 1231.9
    liquidus_c = 1526.8
    melting_range_c = liquidus_c - solidus_c

    doc = Document()

    # Title
    title = doc.add_heading('Cu–Ni Binary System: Phase Diagram and Solidification Analysis', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_paragraph(doc, 'This report analyzes the Cu–Ni binary system using computed phase diagrams and Gibbs energy surfaces. It identifies the solidification start/end temperatures at a representative composition (Cu–40Ni), discusses the phase transformation path during cooling, explains the complete solid solubility, and relates the findings to alloy design and processing.')

    # Section: Phase diagram
    add_heading(doc, '1. Phase Diagram Overview', 1)
    add_paragraph(doc, 'The Cu–Ni system is isomorphous with a single solid solution phase (FCC_A1) across the entire composition range. The calculated binary phase diagram below spans 400–1800 K and shows liquid (L), FCC_A1 (solid solution), and the two-phase (L + FCC_A1) field.')
    add_image(doc, 'CuNi.png', 'Figure 1. Cu–Ni binary phase diagram (computed, 400–1800 K, Δx=0.01).')

    # Section: Solidification begin/end
    add_heading(doc, '2. Solidification Start and End Temperatures (Cu–40Ni)', 1)
    add_paragraph(doc, f'Based on equilibrium calculations for X(Ni)=0.40, cooling from the liquid: '
                       f'solidification begins at the liquidus and ends at the solidus.')
    add_bullet(doc, f'Solidification begins (Liquidus): {liquidus_c:.1f} °C')
    add_bullet(doc, f'Solidification ends (Solidus): {solidus_c:.1f} °C')
    add_bullet(doc, f'Melting/Freezing range: {melting_range_c:.1f} °C')
    add_paragraph(doc, 'Interpretation: On cooling, the first solid (FCC_A1) nucleates at the liquidus; the fraction of solid grows through the two-phase field until the alloy becomes fully solid at the solidus.')

    # Section: Phase transformation path
    add_heading(doc, '3. Phase Transformation Path During Cooling', 1)
    add_paragraph(doc, 'At Cu–40Ni under 1 atm, the expected equilibrium path is:')
    add_bullet(doc, 'Above liquidus: Single-phase Liquid (L).')
    add_bullet(doc, 'Between liquidus and solidus: Two-phase L + FCC_A1 (primary solid solution grows while interdendritic liquid enriches/depletes slightly depending on partitioning).')
    add_bullet(doc, 'Below solidus: Single-phase FCC_A1 (homogeneous solid solution).')
    add_paragraph(doc, 'Because the Cu–Ni system is isomorphous, no intermediate intermetallics or second solid phases appear; only the FCC_A1 solid solution forms from the liquid.')

    # Section: Gibbs energy surfaces
    add_heading(doc, '4. Gibbs Energy Surfaces and Equilibrium', 1)
    add_paragraph(doc, 'The Gibbs energy (G^M) curves for candidate phases at a representative temperature illustrate phase stability. Phase equilibria correspond to common tangents between phases; the lower envelope of the energy curves indicates the stable phase(s).')
    add_image(doc, 'CuNi_energy.png', 'Figure 2. Computed molar Gibbs energy vs composition at 1550 K.')
    add_paragraph(doc, 'At 1550 K (between solidus and liquidus for Cu–40Ni), the L and FCC_A1 energy curves are close, and a common tangent indicates L + FCC_A1 coexistence, consistent with the two-phase field in the phase diagram.')

    # Section: Complete solid solubility
    add_heading(doc, '5. Why the Cu–Ni System Exhibits Complete Solid Solubility', 1)
    add_paragraph(doc, 'The Cu–Ni system satisfies the Hume–Rothery conditions for extensive substitutional solubility:')
    add_bullet(doc, 'Crystal structure: Both Cu and Ni are FCC, eliminating structural mismatch.')
    add_bullet(doc, 'Atomic size: Metallic radii are very similar (Cu ≈ 128 pm, Ni ≈ 124 pm; difference < 15%).')
    add_bullet(doc, 'Valency: Comparable metallic valence; no strong driving force to form compounds.')
    add_bullet(doc, 'Electronegativity: Nearly identical (Cu ≈ 1.90, Ni ≈ 1.91 on the Pauling scale), minimizing chemical ordering tendencies.')
    add_paragraph(doc, 'Together, these factors favor a continuous substitutional solid solution (FCC_A1) over the entire composition range, which is reflected in the isomorphous phase diagram.')

    # Section: Implications for alloy design and processing
    add_heading(doc, '6. Implications for Alloy Design and Processing', 1)
    add_bullet(doc, 'Compositions of interest: 90–10 and 70–30 Cu–Ni alloys are common for marine service due to excellent corrosion resistance and good thermal conductivity.')
    add_bullet(doc, f'Solidification window: For Cu–40Ni, ΔT ≈ {melting_range_c:.0f} °C. A finite freezing range implies dendritic solidification; however, Cu–Ni partitioning is mild (k close to 1), so macrosegregation is limited compared with eutectic systems.')
    add_bullet(doc, 'Homogenization: Post-cast homogenization can remove residual microsegregation efficiently due to single-phase FCC_A1 at service temperatures.')
    add_bullet(doc, 'Weldability: Isomorphous behavior avoids brittle intermetallics; however, a non-zero freezing range can still pose hot cracking risk—control heat input and composition near the diagram center to mitigate.')
    add_bullet(doc, 'Heat treatment: No precipitation hardening; strengthening is via solid solution and cold work. Annealing restores ductility with predictable single-phase recovery/recrystallization behavior.')
    add_bullet(doc, 'Processing maps: The phase diagram enables selection of casting temperatures (above liquidus) and hot-working windows (single-phase FCC_A1), and informs solutionizing temperatures without risking secondary phase formation.')

    # Section: Conclusions
    add_heading(doc, '7. Conclusions', 1)
    add_bullet(doc, f'Solidification of Cu–40Ni begins at {liquidus_c:.1f} °C (liquidus) and ends at {solidus_c:.1f} °C (solidus).')
    add_bullet(doc, 'The cooling path is L → (L + FCC_A1) → FCC_A1.')
    add_bullet(doc, 'Complete solid solubility arises from matched structure, size, valency, and electronegativity.')
    add_bullet(doc, 'These features make Cu–Ni alloys versatile for corrosion-resistant applications and robust processing routes.')

    out_path = 'CuNi_Analysis.docx'
    doc.save(out_path)
    print(f"Report written to {out_path}")

if __name__ == '__main__':
    main()

